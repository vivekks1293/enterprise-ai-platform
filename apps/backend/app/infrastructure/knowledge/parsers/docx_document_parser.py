from collections.abc import AsyncIterator
from io import BytesIO

from docx import Document as DocxDocument

from app.application.knowledge.contracts.parsed_document import (
    ParsedDocument,
)
from app.application.knowledge.contracts.parsed_document_section import (
    ParsedDocumentSection,
)
from app.application.knowledge.ports.document_parser import (
    DocumentParser,
)


class DocxDocumentParser(DocumentParser):
    """
    Extracts textual content from DOCX documents.
    """

    async def parse(
        self,
        *,
        content: bytes,
        filename: str,
    ) -> ParsedDocument:

        # raw_content = bytearray()

        # async for chunk in content:
        #     raw_content.extend(chunk)

        document = DocxDocument(
            BytesIO(content)
        )

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        text = "\n\n".join(paragraphs)

        sections: list[ParsedDocumentSection] = []

        if text:
            sections.append(
                ParsedDocumentSection(
                    content=text,
                    metadata={
                        "source_filename": filename,
                    },
                )
            )

        return ParsedDocument(
            sections=sections,
            metadata={
                "source_filename": filename,
                "parser": "docx",
                "paragraph_count": len(paragraphs),
            },
        )